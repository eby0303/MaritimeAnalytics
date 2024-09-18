import streamlit as st
import pandas as pd
import folium
from streamlit_folium import folium_static
import numpy as np
from datetime import datetime
from docx import Document
import os










# Define DataFrame
np.random.seed(42)

# Define latitude and longitude ranges for the Gulf of Mexico
lat_min, lat_max = 23.0, 28.0
lon_min, lon_max = -95.0, -85.0

# Generate random latitudes and longitudes for vessels
random_latitudes = np.random.uniform(lat_min, lat_max, 12)
random_longitudes = np.random.uniform(lon_min, lon_max, 12)

# Create DataFrame with vessel coordinates and details
data = pd.DataFrame({
    'MMSI': [123456789, 987654321, 111222333, 222333444, 333444555, 444555666, 555666777,
             666777888, 777888999, 888999111, 999111222, 101112131],
    'LATITUDE': random_latitudes,
    'LONGITUDE': random_longitudes,
    'COURSE': [180, 45, 90, 135, 270, 315, 60, 30, 75, 120, 150, 300],
    'SPEED': [12.5, 15.0, 14.0, 16.0, 13.5, 10.0, 11.0, 12.0, 13.0, 14.5, 15.5, 16.5],
    'HEADING': [180, 45, 90, 135, 270, 315, 60, 30, 75, 120, 150, 300],
    'NAVSTAT': [0] * 12,
    'IMO': [9481494, 9231235, 9452541, 9423567, 9293407, 9347934, 9321234, 
            9439875, 9262339, 9235641, 9295876, 9261234],
    'NAME': ['MSC Gülsün', 'Ever Given', 'Maersk Alabama', 'Emma Maersk', 'CMA CGM Marco Polo', 'Hanjin Europe', 
             'Cosco Development', 'MSC Oscar', 'OOCL Hong Kong', 'Madrid Maersk', 'CMA CGM Jules Verne', 'Hyundai Earth'],
    'CALLSIGN': ['3FWI9', 'VRBH9', 'WDE5165', 'OYYH', 'FAYX', '9V9538', 'VRDE8', 'S9IN', 'VRJH4', 'OXVB', 'FMLC', '9V9713'],
    'FLAG': ['Panama', 'Panama', 'United States', 'Denmark', 'France', 'South Korea', 
             'Hong Kong', 'Panama', 'Hong Kong', 'Denmark', 'France', 'South Korea'],
    'AIS Transponder Class': ['Class A'] * 12,
    'General Vessel Type': ['Cargo Ship', 'Cargo Ship', 'Oil Tanker', 'Passenger Ship', 'Bulk Carrier', 
                            'LNG Carrier', 'Ro-Ro Ship', 'Cargo Ship', 'Bulk Carrier', 'Passenger Ship', 
                            'LPG Carrier', 'Container Ship'],
    'Detailed Vessel Type': ['Container Ship', 'Container Ship', 'VLCC (Very Large Crude Carrier)', 
                             'Cruise Ship', 'Capesize Bulk Carrier', 'LNG Tanker', 
                             'Ro-Ro Vessel', 'Feeder Vessel', 'Handysize Bulk Carrier', 
                             'Ferry', 'LPG Tanker', 'Ultra Large Container Vessel'],
    'A': [50 + i*5 for i in range(12)],
    'B': [120 + i*10 for i in range(12)],
    'C': [20 + i*2 for i in range(12)],
    'D': [18 + i*2 for i in range(12)],
    'DRAUGHT': [5.5 + i*0.1 for i in range(12)],
    'DESTINATION': ['Shanghai', 'Rotterdam', 'Newark', 'Singapore', 'Hamburg', 'Busan', 'Jeddah', 'Antwerp', 
                    'Felixstowe', 'Los Angeles', 'Le Havre', 'Kaohsiung'],
    'LOCODE': ['CNSHA', 'NLRTM', 'USNWK', 'SGSIN', 'DEHAM', 'KRPUS', 'SAJED', 'BEANR', 'GBFXT', 
               'USLAX', 'FRLEH', 'TWKHH'],
    'ETA_AIS': [datetime.utcnow().strftime('%Y-%m-%d %H:%M') for _ in range(12)],
    'ETA': [datetime.utcnow().strftime('%Y-%m-%d %H:%M') for _ in range(12)],
    'ETA_PREDICTED': [datetime.utcnow().strftime('%Y-%m-%d %H:%M') for _ in range(12)],
    'DISTANCE_REMAINING': [100 + i*10 for i in range(12)],
    'SRC': ['SAT'] * 12,
    'ZONE': ['Gulf of Mexico'] * 12,
    'ECA': [False] * 12,
    'TIMESTAMP': [datetime.utcnow()] * 12,
    'STATE': ['green'] * 12  # Initial state for all vessels
})



DOCS_DIR = "alerted_documents" 

# Function to create a vessel report as a Word document
def create_vessel_document(vessel_data, imo):
    vessel = vessel_data[vessel_data['IMO'] == imo]
    if vessel.empty:
        return f"Vessel with IMO {imo} not found."

    # Create a new document
    doc = Document()
    doc.add_heading(f"Vessel Report - IMO: {imo}", 0)

    # Add vessel details to the document
    for column in vessel.columns:
        doc.add_paragraph(f"{column}: {vessel[column].values[0]}")

    # Ensure the DOCS_DIR exists
    if not os.path.exists(DOCS_DIR):
        os.makedirs(DOCS_DIR)

    # Save the document with the name "vessel_report_<imo>.docx" in the DOCS_DIR directory
    file_name = os.path.join(DOCS_DIR, f"vessel_report_{imo}.docx")
    doc.save(file_name)

    return f"Document saved as {file_name}"
    
    
# Set the directory where the alerted documents are stored
DOCS_DIR = "alerted_documents" 

def document_viewer():
    # Streamlit app title
    st.title("Alerted Documents Viewer")

    # Function to list all files in the document directory
    def list_documents():
        return [f for f in os.listdir(DOCS_DIR) if f.endswith('.docx') or f.endswith('.pdf')]

    # List all the documents
    documents = list_documents()

    if documents:
        st.subheader("Available Alerted Documents:")
        for doc in documents:
            doc_path = os.path.join(DOCS_DIR, doc)
            # Create a download link for each document
            with open(doc_path, "rb") as file:
                st.download_button(
                    label=f"Download {doc}",
                    data=file,
                    file_name=doc,
                    mime="application/octet-stream"
                )
    else:
        st.warning("No alerted documents available.")

    # Additional info or logging section
    st.write("These documents represent all the alerts generated for anomalies detected in vessel activities.")  
    

    